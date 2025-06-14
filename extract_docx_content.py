#!/usr/bin/env python3
"""
Script to extract text content from a Microsoft Word document (.docx)
"""

import sys
import os

def extract_docx_content(file_path):
    """
    Extract text content from a Word document
    
    Args:
        file_path (str): Path to the Word document
        
    Returns:
        str: Extracted text content
    """
    try:
        # Try to import python-docx
        try:
            from docx import Document
        except ImportError:
            print("Error: python-docx library is not installed.")
            print("Please install it using: pip install python-docx")
            return None
        
        # Check if file exists
        if not os.path.exists(file_path):
            print(f"Error: File '{file_path}' not found.")
            return None
            
        # Open and read the document
        doc = Document(file_path)
        
        # Extract all text from paragraphs
        full_text = []
        
        # Get text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        
        # Get text from tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        
        return '\n'.join(full_text)
        
    except Exception as e:
        print(f"Error reading document: {str(e)}")
        return None

def main():
    # Path to the Word document
    docx_path = "/home/sdhre/transcribe/history till now.docx"
    
    print(f"Extracting content from: {docx_path}")
    print("-" * 80)
    
    # Extract content
    content = extract_docx_content(docx_path)
    
    if content:
        print("Document Content:")
        print("=" * 80)
        print(content)
        print("=" * 80)
        print(f"\nTotal characters: {len(content)}")
        print(f"Total words (approx): {len(content.split())}")
    else:
        print("Failed to extract content from the document.")

if __name__ == "__main__":
    # Path to the Word document
    docx_path = "/home/sdhre/transcribe/history till now.docx"
    
    print(f"Extracting content from: {docx_path}")
    print("-" * 80)
    
    # Extract content
    content = extract_docx_content(docx_path)
    
    if content:
        print("Document Content:")
        print("=" * 80)
        print(content)
        print("=" * 80)
        print(f"\nTotal characters: {len(content)}")
        print(f"Total words (approx): {len(content.split())}")
        
        # Save to a text file for easier access
        output_path = "/home/sdhre/transcribe/history_till_now_extracted.txt"
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            print(f"\nContent saved to: {output_path}")
        except Exception as e:
            print(f"\nError saving to file: {str(e)}")
    else:
        print("Failed to extract content from the document.")